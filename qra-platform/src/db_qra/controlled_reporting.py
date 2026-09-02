# ruff: noqa: E501
"""Controlled stage-6 report context, narrative, validation, and deliverables.

The language model is deliberately kept outside every numerical path.  It may
only return a ``report-draft-v1`` object whose placeholders reference immutable
metrics in ``report-context-v1``.  Rendering, substitution, citations, release
gates, watermarks, and downloadable files remain deterministic.
"""

from __future__ import annotations

import base64
import hashlib
import io
import json
import math
import re
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime, timezone
from html import escape
from pathlib import Path
from typing import Any, Protocol

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from jsonschema import Draft202012Validator
from PIL import Image as PilImage
from PIL import ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as PdfImage,
)
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from .database import QraDatabase, bytes_sha256, json_sha256
from .paths import PROJECT_ROOT

REPORT_SERVICE_VERSION = "qra.controlled-report/1.0.0"
REPORT_CONTEXT_VERSION = "report-context-v1"
REPORT_DRAFT_VERSION = "report-draft-v1"
REPORT_PROMPT_VERSION = "qra-controlled-report-prompt-v1"
REPORT_VALIDATION_VERSION = "controlled-report-validation-v1"
FALLBACK_PROVIDER_ID = "deterministic-template-v1"
SYNTHETIC_WATERMARK = "合成数据 · 仅供软件测试"

STAGE6_ROOT = (
    PROJECT_ROOT / "resources" / "synthetic" / "full-chain-v1" / "stage6"
)
CONTEXT_SCHEMA_PATH = STAGE6_ROOT / "report-context-v1.schema.json"
DRAFT_SCHEMA_PATH = STAGE6_ROOT / "report-draft-v1.schema.json"
PROMPT_PATH = STAGE6_ROOT / "report-prompt-v1.md"
DEMO_ENVELOPE_PATH = (
    PROJECT_ROOT
    / "resources"
    / "synthetic"
    / "full-chain-v1"
    / "stage2"
    / "generated"
    / "S00_BASELINE_D00_CLEAN"
    / "golden"
    / "expected-snapshot.json"
)
DEMO_SOURCE_ROOT = DEMO_ENVELOPE_PATH.parents[1] / "source-documents"

REPORT_SECTIONS = (
    ("project_overview", "项目概述"),
    ("synthetic_boundary", "合成数据和使用边界"),
    ("input_completeness", "输入资料与数据完整度"),
    ("models_parameters", "模型和参数说明"),
    ("failure_frequency", "失效频率和失效概率"),
    ("consequence_analysis", "后果分析"),
    ("individual_risk", "个人风险 IR"),
    ("societal_risk", "社会风险 F-N 和 PLL"),
    ("segment_ranking", "管段风险排序"),
    ("risk_drivers", "风险驱动因素"),
    ("scenario_comparison", "场景对比"),
    ("uncertainty_gaps", "不确定性和数据缺口"),
    ("recommendations", "建议措施"),
    ("formal_release_blockers", "正式发布阻断项"),
    ("reference_directory", "输入、模型、结果和证据引用目录"),
)

PROHIBITED_CLAIMS = (
    "正式评价通过",
    "风险可接受",
    "符合监管要求",
    "正式发布阻断已解除",
    "已解除发布阻断",
    "可直接用于安全决策",
    "完全安全",
)

METRIC_TOKEN = re.compile(r"\{\{metric:([A-Z][A-Z0-9_.-]*)\}\}")
COMPLETION_WORDS = ("已完成", "成功运行", "完整执行", "已经完成")


class StructuredReportProvider(Protocol):
    """Small provider boundary for an online structured-output model."""

    provider_id: str

    def generate(self, *, prompt: str, context: dict[str, Any]) -> dict[str, Any]: ...


@dataclass(frozen=True)
class ControlledReportBuild:
    report: dict[str, Any]
    context: dict[str, Any]
    draft: dict[str, Any]
    validation: dict[str, Any]
    html: bytes
    pdf: bytes
    docx: bytes
    charts: dict[str, bytes]


def _read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _as_float(value: Any, default: float = 0.0) -> float:
    try:
        result = float(value)
    except (TypeError, ValueError):
        return default
    return result if math.isfinite(result) else default


def _scientific(value: Any) -> str:
    return f"{_as_float(value):.6e}"


def _decimal(value: Any, places: int = 3) -> str:
    return f"{_as_float(value):.{places}f}"


def _percent(value: Any, places: int = 2) -> str:
    return f"{100.0 * _as_float(value):.{places}f}%"


def _metric(
    value: Any,
    display: str,
    unit: str | None,
    source_result_ref: str,
) -> dict[str, Any]:
    return {
        "value": value,
        "display": display,
        "unit": unit,
        "source_result_ref": source_result_ref,
    }


def _get(mapping: Any, *path: str, default: Any = None) -> Any:
    value = mapping
    for key in path:
        if not isinstance(value, dict) or key not in value:
            return default
        value = value[key]
    return value


def _image_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    windows = Path("C:/Windows/Fonts")
    candidates = (
        [windows / "msyhbd.ttc", windows / "simhei.ttf", windows / "arialbd.ttf"]
        if bold
        else [windows / "msyh.ttc", windows / "simsun.ttc", windows / "arial.ttf"]
    )
    candidates.extend(
        [
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
            if bold
            else Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        ]
    )
    for path in candidates:
        if path.is_file():
            try:
                return ImageFont.truetype(str(path), size=size, index=0)
            except OSError:
                continue
    return ImageFont.load_default()


def _png_bytes(image: PilImage.Image) -> bytes:
    output = io.BytesIO()
    image.save(output, format="PNG", compress_level=9, optimize=False)
    return output.getvalue()


def _ranking_chart(rows: list[dict[str, Any]]) -> bytes:
    ranked = rows[:10]
    width, height = 1400, 760
    image = PilImage.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title = _image_font(42, bold=True)
    label = _image_font(27, bold=True)
    small = _image_font(23)
    draw.text((70, 42), "Segment PLL ranking", fill="#173F67", font=title)
    draw.text(
        (70, 98),
        "Deterministic values from human_qra · fatalities/year",
        fill="#66788A",
        font=small,
    )
    maximum = max((_as_float(row.get("risk_value_fatalities_per_year")) for row in ranked), default=1.0)
    chart_left, chart_right = 245, 1280
    top, row_height = 155, 54
    for index, row in enumerate(ranked):
        y = top + index * row_height
        value = _as_float(row.get("risk_value_fatalities_per_year"))
        draw.text((70, y + 10), str(row.get("segment_id") or "—"), fill="#243447", font=label)
        draw.rounded_rectangle(
            (chart_left, y + 7, chart_right, y + 41),
            radius=10,
            fill="#E9F0F4",
        )
        bar_right = chart_left + int((chart_right - chart_left) * value / maximum)
        draw.rounded_rectangle(
            (chart_left, y + 7, max(chart_left + 4, bar_right), y + 41),
            radius=10,
            fill="#2F75B5" if index else "#D27A3C",
        )
        draw.text((chart_right - 185, y + 10), f"{value:.3e}", fill="#173F67", font=small)
    draw.line((70, 714, 1330, 714), fill="#D9E2E8", width=2)
    draw.text((70, 722), SYNTHETIC_WATERMARK, fill="#8B9AA4", font=small)
    return _png_bytes(image)


def _fn_chart(points: list[dict[str, Any]]) -> bytes:
    width, height = 1400, 760
    image = PilImage.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title = _image_font(42, bold=True)
    label = _image_font(25)
    draw.text((70, 42), "Societal risk F-N curve", fill="#173F67", font=title)
    left, top, right, bottom = 150, 145, 1290, 650
    draw.rectangle((left, top, right, bottom), outline="#8EA5B2", width=2)
    positive = [
        (
            max(_as_float(point.get("fatalities_at_least")), 1.0e-12),
            max(_as_float(point.get("cumulative_frequency_per_year")), 1.0e-12),
        )
        for point in points
    ]
    if not positive:
        positive = [(1.0, 1.0e-8)]
    log_x = [math.log10(row[0]) for row in positive]
    log_y = [math.log10(row[1]) for row in positive]
    min_x, max_x = min(log_x), max(log_x)
    min_y, max_y = min(log_y), max(log_y)
    if min_x == max_x:
        max_x += 1.0
    if min_y == max_y:
        max_y += 1.0
    min_y -= 0.15
    max_y += 0.15

    def px(value: float) -> int:
        return int(left + (value - min_x) / (max_x - min_x) * (right - left))

    def py(value: float) -> int:
        return int(bottom - (value - min_y) / (max_y - min_y) * (bottom - top))

    for step in range(6):
        y_value = min_y + (max_y - min_y) * step / 5
        y = py(y_value)
        draw.line((left, y, right, y), fill="#E4EBEF", width=2)
        draw.text((42, y - 14), f"{10 ** y_value:.2e}", fill="#66788A", font=label)
    coordinates = [(px(x), py(y)) for x, y in zip(log_x, log_y, strict=True)]
    if len(coordinates) > 1:
        draw.line(coordinates, fill="#D27A3C", width=6, joint="curve")
    for x, y in coordinates:
        draw.ellipse((x - 9, y - 9, x + 9, y + 9), fill="#173F67", outline="white", width=2)
    draw.text((610, 682), "Fatalities N (log scale)", fill="#243447", font=label)
    draw.text((70, 718), SYNTHETIC_WATERMARK, fill="#8B9AA4", font=label)
    return _png_bytes(image)


def _matrix_chart(matrix: dict[str, Any]) -> bytes:
    width, height = 1200, 820
    image = PilImage.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title = _image_font(40, bold=True)
    label = _image_font(25, bold=True)
    small = _image_font(22)
    draw.text((70, 40), "Display risk matrix", fill="#173F67", font=title)
    draw.text(
        (70, 92),
        "Display aid only · not a formal acceptance criterion",
        fill="#9B1C1C",
        font=small,
    )
    cells = {
        (int(row.get("likelihood_grade", 0)), str(row.get("consequence_grade", ""))): row
        for row in matrix.get("cells", [])
        if isinstance(row, dict)
    }
    fills = {
        "LOW": "#EAF3E6",
        "MEDIUM": "#FFE699",
        "MEDIUM_HIGH": "#F4B183",
        "HIGH": "#E98B96",
    }
    left, top, size = 230, 155, 108
    for likelihood in range(5, 0, -1):
        y = top + (5 - likelihood) * size
        draw.text((130, y + 37), f"L{likelihood}", fill="#243447", font=label)
        for consequence_number in range(1, 6):
            consequence = chr(ord("A") + consequence_number - 1)
            x = left + (consequence_number - 1) * size
            cell = cells.get((likelihood, consequence), {})
            band = str(cell.get("display_risk_band") or "LOW")
            draw.rectangle(
                (x, y, x + size, y + size),
                fill=fills.get(band, "#F2F4F7"),
                outline="white",
                width=4,
            )
            count = len(cell.get("segment_ids") or [])
            draw.text((x + 43, y + 25), str(count), fill="#173F67", font=label)
            band_label = {
                "LOW": "LOW",
                "MEDIUM": "MEDIUM",
                "MEDIUM_HIGH": "M-HIGH",
                "HIGH": "HIGH",
            }.get(band, band)
            draw.text((x + 12, y + 68), band_label, fill="#3E4D56", font=small)
    for consequence_number in range(1, 6):
        consequence = chr(ord("A") + consequence_number - 1)
        draw.text((left + (consequence_number - 1) * size + 43, 710), consequence, fill="#243447", font=label)
    draw.text((382, 756), "Consequence grade", fill="#243447", font=small)
    draw.text((70, 786), SYNTHETIC_WATERMARK, fill="#8B9AA4", font=small)
    return _png_bytes(image)


def build_chart_bundle(context_data: dict[str, Any]) -> dict[str, bytes]:
    return {
        "segment_pll_ranking": _ranking_chart(context_data.get("segment_ranking", [])),
        "fn_curve": _fn_chart(context_data.get("fn_curve", [])),
        "risk_matrix": _matrix_chart(context_data.get("risk_matrix", {})),
    }


def _evidence_index(
    database: QraDatabase,
    project: dict[str, Any],
    snapshot: dict[str, Any],
) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    if bool(project.get("is_demo")) and DEMO_SOURCE_ROOT.is_dir():
        for index, path in enumerate(
            sorted(item for item in DEMO_SOURCE_ROOT.iterdir() if item.is_file()), start=1
        ):
            rows.append(
                {
                    "evidence_id": f"EVIDENCE.SOURCE.{index:02d}",
                    "label": path.name,
                    "source": (
                        "synthetic://full-chain-v1/S00_BASELINE/D00_CLEAN/"
                        f"{path.name}"
                    ),
                    "sha256": _sha256_file(path),
                }
            )
    elif project.get("conversion_job_id"):
        for index, source in enumerate(
            database.list_conversion_sources(str(project["conversion_job_id"])), start=1
        ):
            source_hash = str(
                source.get("sha256")
                or source.get("content_sha256")
                or source.get("stored_sha256")
                or json_sha256(source)
            )
            rows.append(
                {
                    "evidence_id": f"EVIDENCE.SOURCE.{index:02d}",
                    "label": str(
                        source.get("relative_path")
                        or source.get("original_file_name")
                        or source.get("id")
                    ),
                    "source": f"conversion-source:{source.get('id')}",
                    "sha256": source_hash,
                }
            )
    if not rows:
        rows.append(
            {
                "evidence_id": "EVIDENCE.SNAPSHOT",
                "label": "不可变输入快照",
                "source": f"input-snapshot:{snapshot['id']}",
                "sha256": str(snapshot["payload_sha256"]),
            }
        )
    review_events = [
        event
        for event in database.list_audit_events(limit=500)
        if event.get("event_type") == "STAGE3_SNAPSHOT_CONFIRMED"
        and str(event.get("entity_id") or "") == str(snapshot["id"])
    ]
    if review_events:
        decision_hash = str((review_events[0].get("detail") or {}).get("decision_set_sha256") or "")
        if re.fullmatch(r"[0-9a-f]{64}", decision_hash):
            rows.append(
                {
                    "evidence_id": "EVIDENCE.REVIEW_DECISION_SET",
                    "label": "人工复核决定集",
                    "source": f"audit-event:{review_events[0]['id']}",
                    "sha256": decision_hash,
                }
            )
    manifest_hash = json_sha256(
        [{"evidence_id": row["evidence_id"], "sha256": row["sha256"]} for row in rows]
    )
    return [
        {
            "evidence_id": "EVIDENCE.SOURCE_PACK_MANIFEST",
            "label": "受控证据索引清单",
            "source": f"input-snapshot:{snapshot['id']}:evidence-index",
            "sha256": manifest_hash,
        },
        *rows,
    ]


def _parameter_bindings(project: dict[str, Any], snapshot_payload: dict[str, Any]) -> list[dict[str, Any]]:
    if bool(project.get("is_demo")) and DEMO_ENVELOPE_PATH.is_file():
        envelope = _read_json(DEMO_ENVELOPE_PATH)
        if json_sha256(envelope.get("qra_input")) == json_sha256(snapshot_payload):
            return list(envelope.get("parameter_pack_bindings") or [])
    bindings: list[dict[str, Any]] = []
    for key in (
        "frequency_library",
        "frequency_correction_model",
        "ignition_model",
        "standard_formula_test_parameters",
    ):
        value = snapshot_payload.get(key)
        if isinstance(value, dict):
            bindings.append(
                {
                    "pack_id": str(
                        value.get("model_id") or value.get("library_id") or key
                    ),
                    "business_content_sha256": json_sha256(value),
                }
            )
    return bindings


def _result_index(
    nodes: list[dict[str, Any]], documents: list[dict[str, Any]]
) -> list[dict[str, Any]]:
    by_node = {str(row["node_id"]): row for row in documents}
    rows = []
    for node in nodes:
        node_id = str(node["node_id"])
        document = by_node.get(node_id)
        rows.append(
            {
                "result_ref": f"RESULT.{node_id}",
                "node_id": node_id,
                "sha256": (
                    str(document["result_sha256"])
                    if document
                    else json_sha256(
                        {
                            "node_id": node_id,
                            "status": node.get("status"),
                            "missing_inputs": node.get("missing_inputs", []),
                        }
                    )
                ),
                "status": str(node.get("status") or "UNKNOWN"),
            }
        )
    return rows


def build_report_context(database: QraDatabase, project_id: str) -> tuple[dict[str, Any], dict[str, bytes]]:
    project = database.get_project(project_id)
    run_id = str(project.get("run_id") or "")
    if not run_id:
        raise ValueError("项目尚未绑定计算任务")
    run = database.get_run(run_id)
    if run.get("status") != "COMPLETED":
        raise ValueError("只有已完成计算才能构建报告上下文")
    snapshot_id = str(run["snapshot_id"])
    snapshot = database.snapshot_metadata(snapshot_id)
    snapshot_payload = database.load_snapshot(snapshot_id)
    if str(run.get("input_sha256")) != str(snapshot.get("payload_sha256")):
        raise ValueError("报告上下文的计算输入哈希与不可变快照不一致")
    nodes = database.list_nodes(run_id)
    if len(nodes) != 11:
        raise ValueError("受控报告上下文必须列出全部11个计算节点")
    documents = database.list_result_documents(run_id)
    results: dict[str, dict[str, Any]] = {}
    for document in documents:
        node_id = str(document["node_id"])
        results[node_id] = database.get_result_document(run_id, node_id)

    geometry = results.get("segment_geometry", {})
    frequency = results.get("failure_frequency", {})
    source_term = results.get("aqt3046_source_term", {})
    jet_fire = results.get("jet_fire_thresholds", {})
    indicator = results.get("indicator_coverage", {})
    human = results.get("human_qra", {})
    matrix = results.get("risk_matrix", {})
    human_risk = _get(human, "human_risk", default={}) or {}
    individual = _get(human_risk, "individual_risk", default={}) or {}
    societal = _get(human_risk, "societal_risk", default={}) or {}
    segment_risk = _get(human_risk, "segment_risk", default={}) or {}
    segment_ranking = list(segment_risk.get("ranking") or [])
    fn_curve = list(societal.get("fn_curve") or [])

    evidence_index = _evidence_index(database, project, snapshot)
    missing_data = []
    for node in nodes:
        for item in node.get("missing_inputs") or []:
            missing_data.append(
                {
                    "node_id": node["node_id"],
                    "path": item.get("path"),
                    "label": item.get("label_zh"),
                    "state": "MISSING_NOT_ZERO",
                }
            )
    for indicator_id in indicator.get("missing_required_indicator_ids") or []:
        missing_data.append(
            {
                "node_id": "indicator_coverage",
                "path": f"engineering_indicators.{indicator_id}",
                "label": str(indicator_id),
                "state": "REQUIRED_INDICATOR_NOT_OBSERVED",
            }
        )

    result_index = _result_index(nodes, documents)
    completed = sum(node.get("status") == "COMPLETED" for node in nodes)
    failed = sum(node.get("status") == "FAILED_ISOLATED" for node in nodes)
    skipped = sum(str(node.get("status")).startswith("SKIPPED") for node in nodes)
    total_frequency = _as_float(frequency.get("total_initiating_frequency_per_year"))
    probability = _as_float(frequency.get("at_least_one_failure_probability_over_horizon"))
    horizon = _as_float(
        _get(frequency, "failure_probability_model", "exposure_years", default=1.0), 1.0
    )
    maximum_ir = _get(individual, "maximum", default={}) or {}
    maximum_mass_flow = max(
        (_as_float(row.get("mass_flow_rate_kg_s")) for row in source_term.get("rows", [])),
        default=0.0,
    )
    maximum_jet_distance = max(
        (_as_float(row.get("threshold_distance_m")) for row in jet_fire.get("rows", [])),
        default=0.0,
    )
    top_segment = segment_ranking[0] if segment_ranking else {}
    mechanisms = frequency.get("frequency_by_mechanism_per_year") or {}
    top_mechanism_id, top_mechanism_value = (
        max(mechanisms.items(), key=lambda item: _as_float(item[1]))
        if mechanisms
        else ("UNKNOWN", 0.0)
    )
    frequency_ranking = frequency.get("segment_ranking") or []
    top_frequency_segment = frequency_ranking[0] if frequency_ranking else {}
    matrix_counts = _get(matrix, "summary", "band_counts", default={}) or {}
    dominant_band, dominant_count = (
        max(matrix_counts.items(), key=lambda item: int(item[1]))
        if matrix_counts
        else ("UNKNOWN", 0)
    )
    synthetic = snapshot_payload.get("synthetic_test_edition") or {}
    parameter_bindings = _parameter_bindings(project, snapshot_payload)

    metrics = {
        "PROJECT.NAME": _metric(project["name"], str(project["name"]), None, "RESULT.data_inventory"),
        "PROJECT.CASE_ID": _metric(
            project.get("case_id") or "—",
            str(project.get("case_id") or "—"),
            None,
            "RESULT.data_inventory",
        ),
        "DATA.CLASSIFICATION": _metric(
            project["data_classification"],
            str(project["data_classification"]),
            None,
            "RESULT.data_inventory",
        ),
        "SNAPSHOT.ID": _metric(snapshot_id, snapshot_id, None, "RESULT.data_inventory"),
        "SNAPSHOT.SHA256": _metric(
            snapshot["payload_sha256"],
            str(snapshot["payload_sha256"]),
            None,
            "RESULT.data_inventory",
        ),
        "RUN.ID": _metric(run_id, run_id, None, "RESULT.data_inventory"),
        "RUN.RESULT_SHA256": _metric(
            run["result_sha256"],
            str(run["result_sha256"]),
            None,
            "RESULT.risk_matrix",
        ),
        "ENGINE.VERSION": _metric(
            run.get("engine_version") or "—",
            str(run.get("engine_version") or "—"),
            None,
            "RESULT.data_inventory",
        ),
        "PARAMETER.PACK_COUNT": _metric(
            len(parameter_bindings), str(len(parameter_bindings)), "packs", "RESULT.data_inventory"
        ),
        "NODE.COMPLETED_COUNT": _metric(
            completed, str(completed), "nodes", "RESULT.data_inventory"
        ),
        "NODE.TOTAL_COUNT": _metric(11, "11", "nodes", "RESULT.data_inventory"),
        "NODE.FAILED_COUNT": _metric(failed, str(failed), "nodes", "RESULT.data_inventory"),
        "NODE.SKIPPED_COUNT": _metric(skipped, str(skipped), "nodes", "RESULT.data_inventory"),
        "DATA.COMPLETENESS_PERCENT": _metric(
            round(completed / 11 * 100),
            f"{round(completed / 11 * 100)}%",
            "%",
            "RESULT.data_inventory",
        ),
        "DATA.EVIDENCE_COUNT": _metric(
            len(evidence_index) - 1,
            str(len(evidence_index) - 1),
            "items",
            "RESULT.data_inventory",
        ),
        "DATA.MISSING_COUNT": _metric(
            len(missing_data), str(len(missing_data)), "items", "RESULT.indicator_coverage"
        ),
        "DATA.INDICATOR_COVERAGE_PERCENT": _metric(
            _as_float(indicator.get("coverage_fraction")) * 100.0,
            _percent(indicator.get("coverage_fraction")),
            "%",
            "RESULT.indicator_coverage",
        ),
        "PIPELINE.SEGMENT_COUNT": _metric(
            int(geometry.get("segment_count") or len(segment_ranking)),
            str(int(geometry.get("segment_count") or len(segment_ranking))),
            "segments",
            "RESULT.segment_geometry",
        ),
        "PIPELINE.TOTAL_LENGTH_KM": _metric(
            _as_float(geometry.get("total_segment_length_km")),
            _decimal(geometry.get("total_segment_length_km"), 3),
            "km",
            "RESULT.segment_geometry",
        ),
        "FAILURE.ANNUAL_FREQUENCY_PER_YEAR": _metric(
            total_frequency,
            _scientific(total_frequency),
            "per_year",
            "RESULT.failure_frequency",
        ),
        "FAILURE.PROBABILITY_HORIZON_YEARS": _metric(
            horizon, _decimal(horizon, 1), "years", "RESULT.failure_frequency"
        ),
        "FAILURE.PROBABILITY_OVER_HORIZON": _metric(
            probability,
            _scientific(probability),
            "probability",
            "RESULT.failure_frequency",
        ),
        "FAILURE.TOP_SEGMENT_ID": _metric(
            top_frequency_segment.get("segment_id") or "—",
            str(top_frequency_segment.get("segment_id") or "—"),
            None,
            "RESULT.failure_frequency",
        ),
        "FAILURE.TOP_MECHANISM_ID": _metric(
            top_mechanism_id, str(top_mechanism_id), None, "RESULT.failure_frequency"
        ),
        "FAILURE.TOP_MECHANISM_FREQUENCY": _metric(
            _as_float(top_mechanism_value),
            _scientific(top_mechanism_value),
            "per_year",
            "RESULT.failure_frequency",
        ),
        "SOURCE.MAX_MASS_FLOW_KG_S": _metric(
            maximum_mass_flow,
            _decimal(maximum_mass_flow, 3),
            "kg/s",
            "RESULT.aqt3046_source_term",
        ),
        "JET_FIRE.MAX_THRESHOLD_DISTANCE_M": _metric(
            maximum_jet_distance,
            _decimal(maximum_jet_distance, 3),
            "m",
            "RESULT.jet_fire_thresholds",
        ),
        "IR.MAX_PER_YEAR": _metric(
            _as_float(maximum_ir.get("value_per_year")),
            _scientific(maximum_ir.get("value_per_year")),
            "per_year",
            "RESULT.human_qra",
        ),
        "IR.MAX_RECEPTOR_ID": _metric(
            maximum_ir.get("cell_id") or "—",
            str(maximum_ir.get("cell_id") or "—"),
            None,
            "RESULT.human_qra",
        ),
        "IR.REFERENCE_LABEL": _metric(
            _get(individual, "acceptability", "label_zh", default="未判定"),
            str(_get(individual, "acceptability", "label_zh", default="未判定")),
            None,
            "RESULT.human_qra",
        ),
        "PLL.PIPELINE_PER_YEAR": _metric(
            _as_float(societal.get("pipeline_pll_per_year")),
            _scientific(societal.get("pipeline_pll_per_year")),
            "fatalities_per_year",
            "RESULT.human_qra",
        ),
        "FN.POINT_COUNT": _metric(len(fn_curve), str(len(fn_curve)), "points", "RESULT.human_qra"),
        "RANK.TOP_SEGMENT_ID": _metric(
            top_segment.get("segment_id") or "—",
            str(top_segment.get("segment_id") or "—"),
            None,
            "RESULT.human_qra",
        ),
        "RANK.TOP_SEGMENT_PLL": _metric(
            _as_float(top_segment.get("risk_value_fatalities_per_year")),
            _scientific(top_segment.get("risk_value_fatalities_per_year")),
            "fatalities_per_year",
            "RESULT.human_qra",
        ),
        "RANK.TOP_SEGMENT_SHARE_PERCENT": _metric(
            _as_float(top_segment.get("fraction_of_pipeline_risk_value")) * 100.0,
            _percent(top_segment.get("fraction_of_pipeline_risk_value")),
            "%",
            "RESULT.human_qra",
        ),
        "RISK_MATRIX.DOMINANT_BAND": _metric(
            dominant_band, str(dominant_band), None, "RESULT.risk_matrix"
        ),
        "RISK_MATRIX.DOMINANT_BAND_COUNT": _metric(
            int(dominant_count), str(int(dominant_count)), "segments", "RESULT.risk_matrix"
        ),
        "SCENARIO.CURRENT_ID": _metric(
            synthetic.get("scenario_id") or project.get("case_id") or "CURRENT",
            str(synthetic.get("scenario_id") or project.get("case_id") or "CURRENT"),
            None,
            "RESULT.data_inventory",
        ),
        "SCENARIO.COMPARISON_COUNT": _metric(1, "1", "scenarios", "RESULT.data_inventory"),
        "FORMAL.ALLOWED": _metric("false", "否", None, "RESULT.human_qra"),
    }

    blockers = list(_get(human, "run", "formal_report_blockers", default=[]) or [])
    if project.get("data_classification") == "SYNTHETIC_TEST_ONLY":
        blockers.insert(0, "全合成数据仅用于软件演示和测试")
    blockers = list(dict.fromkeys(str(item) for item in blockers if str(item).strip()))
    formal_allowed = bool(_get(human, "run", "formal_report_allowed", default=False))
    if project.get("data_classification") == "SYNTHETIC_TEST_ONLY":
        formal_allowed = False

    context_data: dict[str, Any] = {
        "project": {
            "id": str(project["id"]),
            "name": str(project["name"]),
            "case_id": project.get("case_id"),
        },
        "data_classification": str(project["data_classification"]),
        "input_snapshot": {
            "id": snapshot_id,
            "sha256": str(snapshot["payload_sha256"]),
            "schema_version": str(snapshot["schema_version"]),
        },
        "calculation_run": {
            "id": run_id,
            "status": "COMPLETED",
            "result_sha256": str(run["result_sha256"]),
            "engine_version": str(run.get("engine_version") or "unknown"),
        },
        "model_and_parameters": {
            "engine_version": str(run.get("engine_version") or "unknown"),
            "parameter_pack_bindings": parameter_bindings,
        },
        "nodes": [
            {
                "node_id": str(node["node_id"]),
                "sequence_no": int(node["sequence_no"]),
                "label_zh": node.get("label_zh"),
                "status": str(node["status"]),
                "standard_ref": node.get("standard_ref"),
                "missing_inputs": list(node.get("missing_inputs") or []),
            }
            for node in nodes
        ],
        "data_completeness": {
            "percent": round(completed / 11 * 100),
            "completed_node_count": completed,
            "total_node_count": 11,
            "evidence_count": len(evidence_index) - 1,
            "missing_count": len(missing_data),
        },
        "metrics": metrics,
        "segment_ranking": segment_ranking,
        "fn_curve": fn_curve,
        "risk_matrix": matrix,
        "evidence_index": evidence_index,
        "result_index": result_index,
        "missing_data": missing_data,
        "uncertainties": [
            {
                "uncertainty_id": "UNCERTAINTY.SYNTHETIC_BOUNDARY",
                "description": "全部资产、人口、气象、频率和模型参数均为合成测试数据。",
                "source": "EVIDENCE.SOURCE_PACK_MANIFEST",
            },
            {
                "uncertainty_id": "UNCERTAINTY.MODEL_VALIDATION",
                "description": "人员后果与频率修正模型仍保留正式项目验证和批准阻断项。",
                "source": "RESULT.human_qra",
            },
            {
                "uncertainty_id": "UNCERTAINTY.SOCIETAL_CRITERION",
                "description": "当前社会风险曲线用于结果展示，未作为正式项目接受性判据。",
                "source": "RESULT.human_qra",
            },
        ],
        "formal_release_blockers": blockers or ["正式报告条件未满足"],
        "formal_report_allowed": formal_allowed,
    }
    charts = build_chart_bundle(context_data)
    context_data["chart_resources"] = [
        {
            "chart_id": chart_id,
            "path": f"charts/{chart_id}.png",
            "sha256": bytes_sha256(content),
            "content_type": "image/png",
        }
        for chart_id, content in charts.items()
    ]
    basis_hash = json_sha256(context_data)
    context = {
        "schema_version": REPORT_CONTEXT_VERSION,
        "context_id": f"CTX-{basis_hash[:24]}",
        **context_data,
    }
    return context, charts


def _paragraph(
    paragraph_id: str,
    text_template: str,
    *,
    metric_refs: tuple[str, ...] = (),
    evidence_refs: tuple[str, ...] = (),
    result_refs: tuple[str, ...] = (),
    uncertainty_refs: tuple[str, ...] = (),
) -> dict[str, Any]:
    return {
        "paragraph_id": paragraph_id,
        "text_template": text_template,
        "metric_refs": list(metric_refs),
        "evidence_refs": list(evidence_refs),
        "result_refs": list(result_refs),
        "uncertainty_refs": list(uncertainty_refs),
        "prohibited_claim_check": True,
    }


def build_deterministic_draft(context: dict[str, Any]) -> dict[str, Any]:
    """Return a complete, reproducible fifteen-section report draft."""

    sections = [
        {
            "section_id": "project_overview",
            "heading": "项目概述",
            "paragraphs": [
                _paragraph(
                    "project_overview.summary",
                    (
                        "本报告对应项目{{metric:PROJECT.NAME}}，项目编号为"
                        "{{metric:PROJECT.CASE_ID}}。报告使用不可变输入快照"
                        "{{metric:SNAPSHOT.ID}}及计算任务{{metric:RUN.ID}}。"
                    ),
                    metric_refs=(
                        "PROJECT.NAME",
                        "PROJECT.CASE_ID",
                        "SNAPSHOT.ID",
                        "RUN.ID",
                    ),
                    result_refs=("RESULT.data_inventory",),
                )
            ],
        },
        {
            "section_id": "synthetic_boundary",
            "heading": "合成数据和使用边界",
            "paragraphs": [
                _paragraph(
                    "synthetic_boundary.classification",
                    (
                        "本项目数据分类为{{metric:DATA.CLASSIFICATION}}。全部资料、"
                        "参数和结果仅用于软件演示、流程验证与自动化测试，不代表真实"
                        "工程事实，也不得用于真实资产评价、监管申报或工程安全决策。"
                    ),
                    metric_refs=("DATA.CLASSIFICATION",),
                    evidence_refs=("EVIDENCE.SOURCE_PACK_MANIFEST",),
                    uncertainty_refs=("UNCERTAINTY.SYNTHETIC_BOUNDARY",),
                )
            ],
        },
        {
            "section_id": "input_completeness",
            "heading": "输入资料与数据完整度",
            "paragraphs": [
                _paragraph(
                    "input_completeness.summary",
                    (
                        "报告上下文登记{{metric:DATA.EVIDENCE_COUNT}}项源资料证据；"
                        "节点可运行完整度为{{metric:DATA.COMPLETENESS_PERCENT}}，工程"
                        "指标观测覆盖率为{{metric:DATA.INDICATOR_COVERAGE_PERCENT}}。"
                        "当前缺失或未观测项数量为{{metric:DATA.MISSING_COUNT}}。"
                    ),
                    metric_refs=(
                        "DATA.EVIDENCE_COUNT",
                        "DATA.COMPLETENESS_PERCENT",
                        "DATA.INDICATOR_COVERAGE_PERCENT",
                        "DATA.MISSING_COUNT",
                    ),
                    evidence_refs=("EVIDENCE.SOURCE_PACK_MANIFEST",),
                    result_refs=("RESULT.data_inventory", "RESULT.indicator_coverage"),
                )
            ],
        },
        {
            "section_id": "models_parameters",
            "heading": "模型和参数说明",
            "paragraphs": [
                _paragraph(
                    "models_parameters.versions",
                    (
                        "计算引擎版本为{{metric:ENGINE.VERSION}}，本次运行受控绑定"
                        "{{metric:PARAMETER.PACK_COUNT}}个参数包。模型与参数版本仅由"
                        "计算结果和快照绑定关系提供，报告文字不修改模型状态。"
                    ),
                    metric_refs=("ENGINE.VERSION", "PARAMETER.PACK_COUNT"),
                    result_refs=("RESULT.data_inventory", "RESULT.human_qra"),
                    uncertainty_refs=("UNCERTAINTY.MODEL_VALIDATION",),
                )
            ],
        },
        {
            "section_id": "failure_frequency",
            "heading": "失效频率和失效概率",
            "paragraphs": [
                _paragraph(
                    "failure_frequency.summary",
                    (
                        "合成管线年起始失效频率为"
                        "{{metric:FAILURE.ANNUAL_FREQUENCY_PER_YEAR}}，在"
                        "{{metric:FAILURE.PROBABILITY_HORIZON_YEARS}}年时间窗内至少"
                        "发生一次失效的泊松概率为"
                        "{{metric:FAILURE.PROBABILITY_OVER_HORIZON}}。"
                    ),
                    metric_refs=(
                        "FAILURE.ANNUAL_FREQUENCY_PER_YEAR",
                        "FAILURE.PROBABILITY_HORIZON_YEARS",
                        "FAILURE.PROBABILITY_OVER_HORIZON",
                    ),
                    result_refs=("RESULT.failure_frequency",),
                )
            ],
        },
        {
            "section_id": "consequence_analysis",
            "heading": "后果分析",
            "paragraphs": [
                _paragraph(
                    "consequence_analysis.summary",
                    (
                        "确定性源项结果中的最大质量流率为"
                        "{{metric:SOURCE.MAX_MASS_FLOW_KG_S}}，喷射火阈值结果中的最大"
                        "影响距离为{{metric:JET_FIRE.MAX_THRESHOLD_DISTANCE_M}}。这些"
                        "数值由计算节点直接提供，报告生成器不执行重算。"
                    ),
                    metric_refs=(
                        "SOURCE.MAX_MASS_FLOW_KG_S",
                        "JET_FIRE.MAX_THRESHOLD_DISTANCE_M",
                    ),
                    result_refs=(
                        "RESULT.aqt3046_source_term",
                        "RESULT.jet_fire_thresholds",
                    ),
                    uncertainty_refs=("UNCERTAINTY.MODEL_VALIDATION",),
                )
            ],
        },
        {
            "section_id": "individual_risk",
            "heading": "个人风险 IR",
            "paragraphs": [
                _paragraph(
                    "individual_risk.summary",
                    (
                        "最大个人风险为{{metric:IR.MAX_PER_YEAR}}，对应受体"
                        "{{metric:IR.MAX_RECEPTOR_ID}}。系统依据已绑定参考准则显示为"
                        "{{metric:IR.REFERENCE_LABEL}}；该显示不解除本项目的正式发布"
                        "阻断。"
                    ),
                    metric_refs=(
                        "IR.MAX_PER_YEAR",
                        "IR.MAX_RECEPTOR_ID",
                        "IR.REFERENCE_LABEL",
                    ),
                    result_refs=("RESULT.human_qra",),
                    uncertainty_refs=("UNCERTAINTY.MODEL_VALIDATION",),
                )
            ],
        },
        {
            "section_id": "societal_risk",
            "heading": "社会风险 F-N 和 PLL",
            "paragraphs": [
                _paragraph(
                    "societal_risk.summary",
                    (
                        "管线潜在生命损失值为{{metric:PLL.PIPELINE_PER_YEAR}}，"
                        "社会风险曲线包含{{metric:FN.POINT_COUNT}}个受控结果点。当前"
                        "曲线用于软件结果展示，未形成正式项目社会风险接受性判定。"
                    ),
                    metric_refs=("PLL.PIPELINE_PER_YEAR", "FN.POINT_COUNT"),
                    result_refs=("RESULT.human_qra",),
                    uncertainty_refs=("UNCERTAINTY.SOCIETAL_CRITERION",),
                )
            ],
        },
        {
            "section_id": "segment_ranking",
            "heading": "管段风险排序",
            "paragraphs": [
                _paragraph(
                    "segment_ranking.top",
                    (
                        "按管段潜在生命损失值排序，首位管段为"
                        "{{metric:RANK.TOP_SEGMENT_ID}}，其管段潜在生命损失值为"
                        "{{metric:RANK.TOP_SEGMENT_PLL}}，占管线汇总值的"
                        "{{metric:RANK.TOP_SEGMENT_SHARE_PERCENT}}。"
                    ),
                    metric_refs=(
                        "RANK.TOP_SEGMENT_ID",
                        "RANK.TOP_SEGMENT_PLL",
                        "RANK.TOP_SEGMENT_SHARE_PERCENT",
                    ),
                    result_refs=("RESULT.human_qra", "RESULT.risk_matrix"),
                )
            ],
        },
        {
            "section_id": "risk_drivers",
            "heading": "风险驱动因素",
            "paragraphs": [
                _paragraph(
                    "risk_drivers.summary",
                    (
                        "失效频率排序首位管段为{{metric:FAILURE.TOP_SEGMENT_ID}}，"
                        "主导失效机理为{{metric:FAILURE.TOP_MECHANISM_ID}}，该机理年"
                        "频率为{{metric:FAILURE.TOP_MECHANISM_FREQUENCY}}。风险矩阵"
                        "中数量最多的展示色带为{{metric:RISK_MATRIX.DOMINANT_BAND}}，"
                        "包含{{metric:RISK_MATRIX.DOMINANT_BAND_COUNT}}个管段。"
                    ),
                    metric_refs=(
                        "FAILURE.TOP_SEGMENT_ID",
                        "FAILURE.TOP_MECHANISM_ID",
                        "FAILURE.TOP_MECHANISM_FREQUENCY",
                        "RISK_MATRIX.DOMINANT_BAND",
                        "RISK_MATRIX.DOMINANT_BAND_COUNT",
                    ),
                    result_refs=("RESULT.failure_frequency", "RESULT.risk_matrix"),
                )
            ],
        },
        {
            "section_id": "scenario_comparison",
            "heading": "场景对比",
            "paragraphs": [
                _paragraph(
                    "scenario_comparison.scope",
                    (
                        "当前报告上下文的场景为{{metric:SCENARIO.CURRENT_ID}}，纳入"
                        "对比的受控场景数量为{{metric:SCENARIO.COMPARISON_COUNT}}。"
                        "由于本报告只绑定当前计算任务，不对未进入上下文的场景作数值"
                        "推断。"
                    ),
                    metric_refs=(
                        "SCENARIO.CURRENT_ID",
                        "SCENARIO.COMPARISON_COUNT",
                    ),
                    result_refs=("RESULT.data_inventory",),
                )
            ],
        },
        {
            "section_id": "uncertainty_gaps",
            "heading": "不确定性和数据缺口",
            "paragraphs": [
                _paragraph(
                    "uncertainty_gaps.summary",
                    (
                        "数据缺口清单保留{{metric:DATA.MISSING_COUNT}}项。合成数据"
                        "边界、模型验证状态和社会风险判据适用性均作为显式不确定性"
                        "保存，不以默认零值或无来源事实替代。"
                    ),
                    metric_refs=("DATA.MISSING_COUNT",),
                    result_refs=("RESULT.indicator_coverage", "RESULT.human_qra"),
                    uncertainty_refs=(
                        "UNCERTAINTY.SYNTHETIC_BOUNDARY",
                        "UNCERTAINTY.MODEL_VALIDATION",
                        "UNCERTAINTY.SOCIETAL_CRITERION",
                    ),
                )
            ],
        },
        {
            "section_id": "recommendations",
            "heading": "建议措施",
            "paragraphs": [
                _paragraph(
                    "recommendations.controlled",
                    (
                        "后续演示验证应优先复核管段"
                        "{{metric:RANK.TOP_SEGMENT_ID}}的输入证据和主导场景，补充缺失"
                        "项，并在批准的模型与参数版本下重新计算。建议只用于验证软件"
                        "工作流，不构成真实工程整改指令。"
                    ),
                    metric_refs=("RANK.TOP_SEGMENT_ID",),
                    evidence_refs=("EVIDENCE.SOURCE_PACK_MANIFEST",),
                    result_refs=("RESULT.human_qra",),
                    uncertainty_refs=("UNCERTAINTY.SYNTHETIC_BOUNDARY",),
                )
            ],
        },
        {
            "section_id": "formal_release_blockers",
            "heading": "正式发布阻断项",
            "paragraphs": [
                _paragraph(
                    "formal_release_blockers.status",
                    (
                        "正式报告许可为{{metric:FORMAL.ALLOWED}}。本报告持续保留"
                        "合成数据、模型验证、参数批准和判据适用性等阻断项；人工确认"
                        "仅确认测试报告内容，不改变发布许可。"
                    ),
                    metric_refs=("FORMAL.ALLOWED",),
                    result_refs=("RESULT.human_qra",),
                    uncertainty_refs=(
                        "UNCERTAINTY.SYNTHETIC_BOUNDARY",
                        "UNCERTAINTY.MODEL_VALIDATION",
                    ),
                )
            ],
        },
        {
            "section_id": "reference_directory",
            "heading": "输入、模型、结果和证据引用目录",
            "paragraphs": [
                _paragraph(
                    "reference_directory.trace",
                    (
                        "输入快照哈希为{{metric:SNAPSHOT.SHA256}}，计算数值哈希为"
                        "{{metric:RUN.RESULT_SHA256}}。附录中的证据索引、节点结果索引"
                        "和参数绑定目录构成本报告的追溯入口。"
                    ),
                    metric_refs=("SNAPSHOT.SHA256", "RUN.RESULT_SHA256"),
                    evidence_refs=("EVIDENCE.SOURCE_PACK_MANIFEST",),
                    result_refs=("RESULT.data_inventory", "RESULT.risk_matrix"),
                )
            ],
        },
    ]
    assert [row["section_id"] for row in sections] == [row[0] for row in REPORT_SECTIONS]
    return {
        "schema_version": REPORT_DRAFT_VERSION,
        "prompt_version": REPORT_PROMPT_VERSION,
        "generation_mode": "DETERMINISTIC_TEMPLATE_FALLBACK",
        "provider_id": FALLBACK_PROVIDER_ID,
        "sections": sections,
    }


def _schema_check(path: Path, value: dict[str, Any]) -> list[str]:
    validator = Draft202012Validator(_read_json(path))
    return [
        f"{'/'.join(str(part) for part in error.absolute_path) or '$'}: {error.message}"
        for error in sorted(validator.iter_errors(value), key=lambda item: list(item.absolute_path))
    ]


def _check(check_id: str, issues: list[str]) -> dict[str, Any]:
    return {
        "check_id": check_id,
        "status": "PASS" if not issues else "FAIL",
        "issues": issues,
    }


def validate_controlled_report(
    context: dict[str, Any], draft: dict[str, Any]
) -> dict[str, Any]:
    checks: list[dict[str, Any]] = []
    checks.append(_check("CONTEXT_SCHEMA", _schema_check(CONTEXT_SCHEMA_PATH, context)))
    checks.append(_check("DRAFT_SCHEMA", _schema_check(DRAFT_SCHEMA_PATH, draft)))

    section_ids = [row.get("section_id") for row in draft.get("sections", [])]
    expected_ids = [row[0] for row in REPORT_SECTIONS]
    checks.append(
        _check(
            "SECTION_TOPOLOGY",
            [] if section_ids == expected_ids else ["报告章节ID或顺序不符合阶段规范"],
        )
    )

    metric_ids = set(context.get("metrics", {}))
    evidence_ids = {
        str(row.get("evidence_id")) for row in context.get("evidence_index", [])
    }
    result_rows = {
        str(row.get("result_ref")): row for row in context.get("result_index", [])
    }
    uncertainty_ids = {
        str(row.get("uncertainty_id")) for row in context.get("uncertainties", [])
    }
    metric_issues: list[str] = []
    evidence_issues: list[str] = []
    result_issues: list[str] = []
    uncertainty_issues: list[str] = []
    node_issues: list[str] = []
    prohibited_issues: list[str] = []
    conclusion_issues: list[str] = []
    used_metrics: set[str] = set()
    used_evidence: set[str] = set()
    used_results: set[str] = set()

    for section in draft.get("sections", []):
        for paragraph in section.get("paragraphs", []):
            paragraph_id = str(paragraph.get("paragraph_id") or "unknown")
            template = str(paragraph.get("text_template") or "")
            token_refs = METRIC_TOKEN.findall(template)
            declared_metrics = [str(item) for item in paragraph.get("metric_refs", [])]
            used_metrics.update(declared_metrics)
            if set(token_refs) != set(declared_metrics) or len(token_refs) != len(
                declared_metrics
            ):
                metric_issues.append(f"{paragraph_id}: 指标占位符与metric_refs不一致")
            for ref in declared_metrics:
                if ref not in metric_ids:
                    metric_issues.append(f"{paragraph_id}: 未知指标引用 {ref}")
            without_tokens = METRIC_TOKEN.sub("", template)
            if re.search(r"[0-9]", without_tokens):
                metric_issues.append(f"{paragraph_id}: 自由文本包含未受控数字")
            if "{{" in without_tokens or "}}" in without_tokens:
                metric_issues.append(f"{paragraph_id}: 存在未知占位符")

            paragraph_evidence = [
                str(item) for item in paragraph.get("evidence_refs", [])
            ]
            paragraph_results = [str(item) for item in paragraph.get("result_refs", [])]
            paragraph_uncertainties = [
                str(item) for item in paragraph.get("uncertainty_refs", [])
            ]
            used_evidence.update(paragraph_evidence)
            used_results.update(paragraph_results)
            for ref in paragraph_evidence:
                if ref not in evidence_ids:
                    evidence_issues.append(f"{paragraph_id}: 未知证据引用 {ref}")
            for ref in paragraph_results:
                if ref not in result_rows:
                    result_issues.append(f"{paragraph_id}: 未知结果引用 {ref}")
                elif result_rows[ref].get("status") != "COMPLETED" and any(
                    word in template for word in COMPLETION_WORDS
                ):
                    node_issues.append(f"{paragraph_id}: 将未完成节点描述为已完成")
            for ref in paragraph_uncertainties:
                if ref not in uncertainty_ids:
                    uncertainty_issues.append(f"{paragraph_id}: 未知不确定性引用 {ref}")
            if not (
                declared_metrics
                or paragraph_evidence
                or paragraph_results
                or paragraph_uncertainties
            ):
                conclusion_issues.append(f"{paragraph_id}: 段落没有任何受控引用")
            for phrase in PROHIBITED_CLAIMS:
                if phrase in template:
                    prohibited_issues.append(f"{paragraph_id}: 包含禁用表述“{phrase}”")
            if paragraph.get("prohibited_claim_check") is not True:
                prohibited_issues.append(f"{paragraph_id}: 禁用表述检查未声明通过")

    checks.append(_check("NUMERIC_REFERENCES", metric_issues))
    checks.append(_check("EVIDENCE_REFERENCES", evidence_issues))
    checks.append(_check("RESULT_REFERENCES", result_issues))
    checks.append(_check("UNCERTAINTY_REFERENCES", uncertainty_issues))
    checks.append(_check("NODE_STATUS_CONSISTENCY", node_issues))
    checks.append(_check("PROHIBITED_CLAIMS", prohibited_issues))
    checks.append(_check("KEY_CONCLUSION_REFERENCES", conclusion_issues))

    formal_issues = []
    if context.get("data_classification") == "SYNTHETIC_TEST_ONLY":
        if context.get("formal_report_allowed") is not False:
            formal_issues.append("全合成报告的正式报告许可必须为false")
        if not context.get("formal_release_blockers"):
            formal_issues.append("全合成报告必须保留正式发布阻断项")
    checks.append(_check("FORMAL_RELEASE_GATE", formal_issues))

    reference_targets = {
        "metrics": sorted(used_metrics),
        "evidence": sorted(used_evidence),
        "results": sorted(used_results),
    }
    status = "PASS" if all(row["status"] == "PASS" for row in checks) else "FAIL"
    return {
        "schema_version": REPORT_VALIDATION_VERSION,
        "status": status,
        "checks": checks,
        "reference_targets": reference_targets,
        "reference_targets_sha256": json_sha256(reference_targets),
        "context_sha256": json_sha256(context),
        "draft_sha256": json_sha256(draft),
    }


def render_text_template(template: str, context: dict[str, Any]) -> str:
    metrics = context.get("metrics", {})

    def substitute(match: re.Match[str]) -> str:
        metric_id = match.group(1)
        if metric_id not in metrics:
            raise ValueError(f"未知指标引用：{metric_id}")
        return str(metrics[metric_id]["display"])

    rendered = METRIC_TOKEN.sub(substitute, template)
    if "{{" in rendered or "}}" in rendered:
        raise ValueError("报告段落存在未解析占位符")
    return rendered


def _rendered_paragraphs(
    context: dict[str, Any], draft: dict[str, Any]
) -> list[tuple[dict[str, Any], list[str]]]:
    return [
        (
            section,
            [
                render_text_template(str(paragraph["text_template"]), context)
                for paragraph in section["paragraphs"]
            ],
        )
        for section in draft["sections"]
    ]


def _chart_data_uri(content: bytes) -> str:
    return "data:image/png;base64," + base64.b64encode(content).decode("ascii")


def render_controlled_html(
    context: dict[str, Any],
    draft: dict[str, Any],
    validation: dict[str, Any],
    charts: dict[str, bytes],
) -> bytes:
    rendered_sections = _rendered_paragraphs(context, draft)
    metrics = context["metrics"]
    context_hash = json_sha256(context)
    draft_hash = json_sha256(draft)
    sections_html = []
    for index, (section, paragraphs) in enumerate(rendered_sections, start=1):
        paragraphs_html = "".join(f"<p>{escape(text)}</p>" for text in paragraphs)
        visual = ""
        if section["section_id"] == "societal_risk":
            visual = (
                f'<figure><img src="{_chart_data_uri(charts["fn_curve"])}" '
                'alt="F-N曲线"><figcaption>图：社会风险 F-N 曲线</figcaption></figure>'
            )
        elif section["section_id"] == "segment_ranking":
            visual = (
                f'<figure><img src="{_chart_data_uri(charts["segment_pll_ranking"])}" '
                'alt="管段PLL排序"><figcaption>图：管段 PLL 排序</figcaption></figure>'
            )
        elif section["section_id"] == "risk_drivers":
            visual = (
                f'<figure><img src="{_chart_data_uri(charts["risk_matrix"])}" '
                'alt="风险矩阵"><figcaption>图：风险矩阵展示</figcaption></figure>'
            )
        sections_html.append(
            f'<section id="{escape(section["section_id"])}">'
            f"<h2><span>{index:02d}</span>{escape(section['heading'])}</h2>"
            f"{paragraphs_html}{visual}</section>"
        )

    node_rows = "".join(
        "<tr>"
        f"<td>{int(row['sequence_no'])}</td>"
        f"<td>{escape(str(row['label_zh'] or row['node_id']))}</td>"
        f"<td><span class=\"status\">{escape(str(row['status']))}</span></td>"
        f"<td>{escape(str(row['standard_ref'] or '—'))}</td>"
        "</tr>"
        for row in context["nodes"]
    )
    ranking_rows = "".join(
        "<tr>"
        f"<td>{int(row.get('risk_value_rank') or index)}</td>"
        f"<td>{escape(str(row.get('segment_id') or '—'))}</td>"
        f"<td>{escape(_scientific(row.get('risk_value_fatalities_per_year')))}</td>"
        f"<td>{escape(_percent(row.get('fraction_of_pipeline_risk_value')))}</td>"
        "</tr>"
        for index, row in enumerate(context["segment_ranking"][:10], start=1)
    )
    fn_rows = "".join(
        "<tr>"
        f"<td>{escape(_decimal(row.get('fatalities_at_least'), 1))}</td>"
        f"<td>{escape(_scientific(row.get('cumulative_frequency_per_year')))}</td>"
        "</tr>"
        for row in context["fn_curve"]
    )
    evidence_rows = "".join(
        "<tr>"
        f"<td>{escape(str(row['evidence_id']))}</td>"
        f"<td>{escape(str(row['label']))}</td>"
        f"<td><code>{escape(str(row['sha256']))}</code></td>"
        "</tr>"
        for row in context["evidence_index"]
    )
    result_rows = "".join(
        "<tr>"
        f"<td>{escape(str(row['result_ref']))}</td>"
        f"<td>{escape(str(row['status']))}</td>"
        f"<td><code>{escape(str(row['sha256']))}</code></td>"
        "</tr>"
        for row in context["result_index"]
    )
    blocker_items = "".join(
        f"<li>{escape(str(blocker))}</li>" for blocker in context["formal_release_blockers"]
    )
    html = f"""<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{escape(str(context['project']['name']))} · 受控测试报告</title>
<style>
:root{{--navy:#173f67;--blue:#2f75b5;--ink:#243447;--muted:#66788a;--line:#d9e2e8;--soft:#f4f7f9;--warn:#9b1c1c;--orange:#d27a3c}}
*{{box-sizing:border-box}}body{{margin:0;background:#edf2f4;color:var(--ink);font:15px/1.75 "Microsoft YaHei","Noto Sans CJK SC",Arial,sans-serif}}
body::before{{content:"{escape(SYNTHETIC_WATERMARK)}";position:fixed;inset:43% auto auto 18%;transform:rotate(-28deg);font-size:68px;font-weight:800;color:rgba(23,63,103,.055);pointer-events:none;z-index:0;white-space:nowrap}}
main{{position:relative;z-index:1;width:min(1080px,calc(100% - 30px));margin:28px auto;background:white;box-shadow:0 18px 50px rgba(17,47,64,.12)}}
.cover{{min-height:740px;padding:70px 76px;background:linear-gradient(145deg,#102f4c,#1f5a7c 65%,#2f75b5);color:white;display:flex;flex-direction:column}}
.kicker{{letter-spacing:.19em;text-transform:uppercase;color:#b9d6e7;font-size:12px;font-weight:800}}h1{{font-size:42px;line-height:1.18;margin:120px 0 18px;max-width:760px}}
.subtitle{{font-size:18px;color:#d8e8f1;max-width:760px}}.cover-grid{{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:18px;margin-top:auto;border-top:1px solid rgba(255,255,255,.26);padding-top:26px}}
.cover-grid span{{display:block;color:#a9c8d9;font-size:11px}}.cover-grid b{{display:block;overflow-wrap:anywhere}}.watermark{{margin-top:22px;color:#ffd8c0;font-weight:800}}
.content{{padding:52px 76px}}.boundary{{border-left:6px solid var(--warn);background:#fff3f3;padding:17px 20px;margin-bottom:35px;color:#7d2727;font-weight:700}}
.metric-grid{{display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:12px;margin:26px 0 42px}}.metric{{background:var(--soft);border-radius:12px;padding:16px}}
.metric span{{display:block;color:var(--muted);font-size:11px}}.metric b{{display:block;color:var(--navy);font-size:18px;overflow-wrap:anywhere}}
section{{margin:0 0 46px;break-inside:avoid}}h2{{font-size:22px;color:var(--navy);border-bottom:1px solid var(--line);padding-bottom:8px;margin:0 0 16px}}h2 span{{color:var(--orange);font-size:12px;margin-right:12px}}
p{{margin:0 0 12px}}figure{{margin:24px 0}}figure img{{display:block;width:100%;border:1px solid var(--line);border-radius:12px}}figcaption{{text-align:center;color:var(--muted);font-size:12px;margin-top:7px}}
h3{{color:var(--navy);margin:35px 0 10px}}table{{width:100%;border-collapse:collapse;margin:12px 0 28px;font-size:12px}}th,td{{padding:9px 10px;border:1px solid var(--line);vertical-align:top}}th{{background:#e8eef5;color:var(--navy);text-align:left}}code{{font:10px ui-monospace,Consolas,monospace;overflow-wrap:anywhere}}.status{{font-weight:800;color:#23633d}}
.blockers{{background:#fff5ee;border:1px solid #f1cdb5;border-radius:12px;padding:18px 22px}}.audit{{background:#102f4c;color:#dce9f0;padding:28px 76px;font-size:11px}}.audit code{{color:white}}
@media(max-width:700px){{main{{width:100%;margin:0}}.cover,.content,.audit{{padding:32px 22px}}h1{{font-size:31px;margin-top:80px}}.metric-grid,.cover-grid{{grid-template-columns:1fr 1fr}}body::before{{font-size:40px;left:5%}}}}
@media print{{body{{background:white}}main{{width:100%;margin:0;box-shadow:none}}.cover{{page-break-after:always}}section{{break-inside:auto}}}}
</style></head><body><main>
<article class="cover"><div class="kicker">Controlled QRA test report</div><h1>{escape(str(context['project']['name']))}</h1><div class="subtitle">全合成原始资料到完整风险计算的受控报告草稿</div>
<div class="cover-grid"><div><span>项目编号</span><b>{escape(str(context['project']['case_id'] or '—'))}</b></div><div><span>报告上下文</span><b>{escape(str(context['context_id']))}</b></div><div><span>输入快照</span><b>{escape(str(context['input_snapshot']['id']))}</b></div><div><span>计算任务</span><b>{escape(str(context['calculation_run']['id']))}</b></div></div><div class="watermark">{escape(SYNTHETIC_WATERMARK)} · 禁止用于真实资产评价</div></article>
<div class="content"><div class="boundary">本报告由受控结构化草稿生成。所有数字均来自不可变计算结果；正式报告许可保持关闭。</div>
<div class="metric-grid"><div class="metric"><span>完成节点</span><b>{escape(metrics['NODE.COMPLETED_COUNT']['display'])} / {escape(metrics['NODE.TOTAL_COUNT']['display'])}</b></div><div class="metric"><span>管线 PLL</span><b>{escape(metrics['PLL.PIPELINE_PER_YEAR']['display'])}</b></div><div class="metric"><span>最大 IR</span><b>{escape(metrics['IR.MAX_PER_YEAR']['display'])}</b></div><div class="metric"><span>正式报告许可</span><b>{escape(metrics['FORMAL.ALLOWED']['display'])}</b></div></div>
{''.join(sections_html)}
<h3>节点状态目录</h3><table><thead><tr><th>序号</th><th>节点</th><th>状态</th><th>依据</th></tr></thead><tbody>{node_rows}</tbody></table>
<h3>管段风险排序（前十）</h3><table><thead><tr><th>排序</th><th>管段</th><th>PLL</th><th>占比</th></tr></thead><tbody>{ranking_rows}</tbody></table>
<h3>F-N 结果点</h3><table><thead><tr><th>死亡人数下限</th><th>累积频率</th></tr></thead><tbody>{fn_rows}</tbody></table>
<h3>正式发布阻断项</h3><div class="blockers"><ul>{blocker_items}</ul></div>
<h3>证据索引</h3><table><thead><tr><th>证据ID</th><th>名称</th><th>SHA-256</th></tr></thead><tbody>{evidence_rows}</tbody></table>
<h3>结果索引</h3><table><thead><tr><th>结果引用</th><th>状态</th><th>SHA-256</th></tr></thead><tbody>{result_rows}</tbody></table>
</div><footer class="audit">上下文哈希 <code>{context_hash}</code><br>草稿哈希 <code>{draft_hash}</code><br>校验状态 <b>{escape(str(validation['status']))}</b> · {escape(SYNTHETIC_WATERMARK)}</footer>
</main></body></html>"""
    return html.encode("utf-8")


def _register_pdf_font() -> str:
    font_name = "QRA-CJK"
    if font_name in pdfmetrics.getRegisteredFontNames():
        return font_name
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for path in candidates:
        if path.is_file():
            try:
                pdfmetrics.registerFont(TTFont(font_name, str(path), subfontIndex=0))
                return font_name
            except Exception:
                continue
    fallback = "STSong-Light"
    if fallback not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont(fallback))
    return fallback


def _pdf_table(
    data: list[list[str]],
    widths: list[float],
    *,
    font_name: str,
    header: bool = True,
) -> Table:
    table = Table(data, colWidths=widths, repeatRows=1 if header else 0, hAlign="LEFT")
    commands: list[tuple[Any, ...]] = [
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#243447")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CAD7DE")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    if header:
        commands.extend(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#173F67")),
                ("FONTSIZE", (0, 0), (-1, 0), 8.0),
            ]
        )
    table.setStyle(TableStyle(commands))
    return table


def render_controlled_pdf(
    context: dict[str, Any],
    draft: dict[str, Any],
    validation: dict[str, Any],
    charts: dict[str, bytes],
) -> bytes:
    output = io.BytesIO()
    font_name = _register_pdf_font()
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "QraTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=28,
        leading=34,
        textColor=colors.HexColor("#173F67"),
        alignment=TA_CENTER,
        spaceAfter=16,
    )
    subtitle = ParagraphStyle(
        "QraSubtitle",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=12,
        leading=18,
        textColor=colors.HexColor("#66788A"),
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    h1 = ParagraphStyle(
        "QraH1",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=15,
        leading=20,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=18,
        spaceAfter=10,
    )
    h2 = ParagraphStyle(
        "QraH2",
        parent=h1,
        fontSize=12,
        leading=16,
        spaceBefore=14,
        spaceAfter=7,
    )
    body = ParagraphStyle(
        "QraBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=9.5,
        leading=15,
        textColor=colors.HexColor("#243447"),
        alignment=TA_LEFT,
        spaceAfter=7,
    )
    note = ParagraphStyle(
        "QraNote",
        parent=body,
        fontSize=8.5,
        leading=13,
        textColor=colors.HexColor("#9B1C1C"),
        borderColor=colors.HexColor("#E5B2B2"),
        borderWidth=0.8,
        borderPadding=9,
        backColor=colors.HexColor("#FFF3F3"),
        spaceAfter=14,
    )
    document = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=f"{context['project']['name']} · 受控测试报告",
        author="QRA Platform",
        subject=SYNTHETIC_WATERMARK,
    )

    def page(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        page_width, page_height = letter
        canvas.setFont(font_name, 8)
        canvas.setFillColor(colors.HexColor("#66788A"))
        canvas.drawString(inch, 0.52 * inch, "QRA 受控测试报告")
        canvas.drawRightString(page_width - inch, 0.52 * inch, f"第 {doc.page} 页")
        canvas.setFillColor(colors.Color(0.09, 0.25, 0.40, alpha=0.055))
        canvas.translate(page_width / 2, page_height / 2)
        canvas.rotate(32)
        canvas.setFont(font_name, 32)
        canvas.drawCentredString(0, 0, SYNTHETIC_WATERMARK)
        canvas.restoreState()

    story: list[Any] = [
        Spacer(1, 0.95 * inch),
        Paragraph(escape(str(context["project"]["name"])), title),
        Paragraph("全合成原始资料到完整风险计算的受控报告草稿", subtitle),
        Spacer(1, 0.35 * inch),
        _pdf_table(
            [
                ["项目编号", str(context["project"]["case_id"] or "—")],
                ["报告上下文", str(context["context_id"])],
                ["输入快照", str(context["input_snapshot"]["id"])],
                ["计算任务", str(context["calculation_run"]["id"])],
                ["数据分类", str(context["data_classification"])],
            ],
            [1.35 * inch, 5.15 * inch],
            font_name=font_name,
            header=False,
        ),
        Spacer(1, 0.45 * inch),
        Paragraph(SYNTHETIC_WATERMARK, note),
        Paragraph(
            "全部资料、参数和结果均为人工合成，不得用于真实资产评价、监管申报或工程安全决策。",
            body,
        ),
        PageBreak(),
    ]
    visual_for_section = {
        "societal_risk": "fn_curve",
        "segment_ranking": "segment_pll_ranking",
        "risk_drivers": "risk_matrix",
    }
    for index, (section, paragraphs) in enumerate(_rendered_paragraphs(context, draft), start=1):
        block: list[Any] = [Paragraph(f"{index:02d}  {escape(section['heading'])}", h1)]
        block.extend(Paragraph(escape(text), body) for text in paragraphs)
        chart_id = visual_for_section.get(section["section_id"])
        if chart_id:
            chart = PdfImage(io.BytesIO(charts[chart_id]), width=6.25 * inch, height=3.4 * inch)
            chart.hAlign = "CENTER"
            block.extend([Spacer(1, 6), chart])
        story.append(KeepTogether(block) if chart_id else block[0])
        if not chart_id:
            story.extend(block[1:])

    story.extend(
        [
            Paragraph("节点状态目录", h2),
            _pdf_table(
                [
                    ["序号", "节点", "状态", "依据"],
                    *[
                        [
                            str(row["sequence_no"]),
                            str(row["label_zh"] or row["node_id"]),
                            str(row["status"]),
                            str(row["standard_ref"] or "—"),
                        ]
                        for row in context["nodes"]
                    ],
                ],
                [0.45 * inch, 1.95 * inch, 1.1 * inch, 3.0 * inch],
                font_name=font_name,
            ),
            KeepTogether(
                [
                    Paragraph("管段风险排序（前十）", h2),
                    _pdf_table(
                        [
                            ["排序", "管段", "PLL", "占比"],
                            *[
                                [
                                    str(row.get("risk_value_rank") or index),
                                    str(row.get("segment_id") or "—"),
                                    _scientific(
                                        row.get("risk_value_fatalities_per_year")
                                    ),
                                    _percent(
                                        row.get("fraction_of_pipeline_risk_value")
                                    ),
                                ]
                                for index, row in enumerate(
                                    context["segment_ranking"][:10], start=1
                                )
                            ],
                        ],
                        [0.7 * inch, 1.4 * inch, 2.4 * inch, 2.0 * inch],
                        font_name=font_name,
                    ),
                ]
            ),
            Paragraph("F-N 结果点", h2),
            _pdf_table(
                [
                    ["死亡人数下限", "累积频率"],
                    *[
                        [
                            _decimal(row.get("fatalities_at_least"), 1),
                            _scientific(row.get("cumulative_frequency_per_year")),
                        ]
                        for row in context["fn_curve"]
                    ],
                ],
                [2.5 * inch, 4.0 * inch],
                font_name=font_name,
            ),
        ]
    )
    release_block = [Paragraph("正式发布阻断项", h2)]
    release_block.extend(
        Paragraph(f"• {escape(str(blocker))}", body)
        for blocker in context["formal_release_blockers"]
    )
    release_block.extend(
        (
            Paragraph("追溯哈希", h2),
            Paragraph(f"上下文哈希：{json_sha256(context)}", body),
            Paragraph(f"草稿哈希：{json_sha256(draft)}", body),
            Paragraph(f"校验状态：{escape(str(validation['status']))}", body),
        )
    )
    story.append(KeepTogether(release_block))
    document.build(story, onFirstPage=page, onLaterPages=page)
    return output.getvalue()


def _set_docx_font(run: Any, *, name: str = "Calibri", east_asia: str = "Microsoft YaHei") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def _set_cell_margins(cell: Any, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    table_width = table_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_pr.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:w"), str(widths_dxa[index]))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _shade_cell(cell: Any, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _format_docx_table(table: Any, *, header: bool = True) -> None:
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        if header and row_index == 0:
            header_property = OxmlElement("w:tblHeader")
            header_property.set(qn("w:val"), "true")
            row._tr.get_or_add_trPr().append(header_property)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if header and row_index == 0:
                _shade_cell(cell, "E8EEF5")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    _set_docx_font(run)
                    run.font.size = Pt(8.5)
                    if header and row_index == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(23, 63, 103)


def _docx_table(document: Document, rows: list[list[str]], widths_dxa: list[int]) -> Any:
    table = document.add_table(rows=0, cols=len(widths_dxa))
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    _set_table_geometry(table, widths_dxa)
    _format_docx_table(table)
    return table


def _append_docx_watermark(document: Document) -> None:
    header = document.sections[0].header
    paragraph = header.add_paragraph()
    watermark = parse_xml(
        f"""<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:v="urn:schemas-microsoft-com:vml"
        xmlns:o="urn:schemas-microsoft-com:office:office"><w:pict>
        <v:shape id="QraSyntheticWatermark"
        o:spid="_x0000_s1025" type="#_x0000_t136"
        style="position:absolute;margin-left:0;margin-top:0;width:468pt;height:468pt;
        rotation:315;z-index:-251654144;mso-position-horizontal:center;
        mso-position-vertical:center;mso-wrap-edited:f" fillcolor="#AEBCC5" stroked="f">
        <v:fill opacity="0.15"/><v:textpath style="font-family:'Calibri';font-size:1pt"
        string="{SYNTHETIC_WATERMARK}"/><v:path textpathok="t"/></v:shape></w:pict></w:r>"""
    )
    paragraph._p.append(watermark)


def _configure_docx_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    tokens = {
        "Normal": (11, "243447", 0, 6, 1.25),
        "Heading 1": (16, "2E74B5", 18, 10, 1.0),
        "Heading 2": (13, "2E74B5", 14, 7, 1.0),
        "Heading 3": (12, "1F4D78", 10, 5, 1.0),
    }
    for style_name, (size, color, before, after, line_spacing) in tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.keep_with_next = style_name != "Normal"


def _add_docx_picture(run: Any, image_bytes: bytes, *, width: Any, alt_text: str) -> Any:
    """Add an inline image with the accessibility metadata Word exposes as alt text."""

    shape = run.add_picture(io.BytesIO(image_bytes), width=width)
    doc_properties = shape._inline.docPr
    doc_properties.set("descr", alt_text)
    doc_properties.set("title", alt_text)
    return shape


def render_controlled_docx(
    context: dict[str, Any],
    draft: dict[str, Any],
    validation: dict[str, Any],
    charts: dict[str, bytes],
) -> bytes:
    document = Document()
    _configure_docx_styles(document)
    document.core_properties.title = f"{context['project']['name']} · 受控测试报告"
    document.core_properties.subject = SYNTHETIC_WATERMARK
    document.core_properties.author = "QRA Platform"
    fixed_time = datetime(2026, 9, 1, tzinfo=timezone.utc)
    document.core_properties.created = fixed_time
    document.core_properties.modified = fixed_time

    header = document.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("QRA 受控测试报告  |  仅供软件演示与测试")
    _set_docx_font(header_run)
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor(102, 120, 138)
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run(SYNTHETIC_WATERMARK)
    _set_docx_font(footer_run)
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(155, 28, 28)
    _append_docx_watermark(document)

    for _ in range(5):
        document.add_paragraph()
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    kicker_run = kicker.add_run("CONTROLLED QRA TEST REPORT")
    _set_docx_font(kicker_run)
    kicker_run.bold = True
    kicker_run.font.size = Pt(10.5)
    kicker_run.font.color.rgb = RGBColor(210, 122, 60)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run(str(context["project"]["name"]))
    _set_docx_font(title_run)
    title_run.bold = True
    title_run.font.size = Pt(30)
    title_run.font.color.rgb = RGBColor(32, 55, 72)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    subtitle_run = subtitle.add_run("全合成原始资料到完整风险计算的受控报告草稿")
    _set_docx_font(subtitle_run)
    subtitle_run.font.size = Pt(15)
    subtitle_run.font.color.rgb = RGBColor(43, 81, 99)
    _docx_table(
        document,
        [
            ["字段", "受控值"],
            ["项目编号", str(context["project"]["case_id"] or "—")],
            ["报告上下文", str(context["context_id"])],
            ["输入快照", str(context["input_snapshot"]["id"])],
            ["计算任务", str(context["calculation_run"]["id"])],
            ["数据分类", str(context["data_classification"])],
        ],
        [2700, 6660],
    )
    warning = document.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning.paragraph_format.space_before = Pt(28)
    warning.paragraph_format.space_after = Pt(8)
    warning_run = warning.add_run(SYNTHETIC_WATERMARK)
    _set_docx_font(warning_run)
    warning_run.bold = True
    warning_run.font.size = Pt(13)
    warning_run.font.color.rgb = RGBColor(155, 28, 28)
    boundary = document.add_paragraph()
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    boundary_run = boundary.add_run(
        "全部资料、参数和结果均为人工合成，不得用于真实资产评价、监管申报或工程安全决策。"
    )
    _set_docx_font(boundary_run)
    boundary_run.font.size = Pt(10)
    document.add_page_break()

    visual_for_section = {
        "societal_risk": "fn_curve",
        "segment_ranking": "segment_pll_ranking",
        "risk_drivers": "risk_matrix",
    }
    for index, (section, paragraphs) in enumerate(_rendered_paragraphs(context, draft), start=1):
        heading = document.add_heading(f"{index:02d}  {section['heading']}", level=1)
        heading.paragraph_format.keep_with_next = True
        for text in paragraphs:
            paragraph = document.add_paragraph(text)
            for run in paragraph.runs:
                _set_docx_font(run)
        chart_id = visual_for_section.get(section["section_id"])
        if chart_id:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_together = True
            alt_text = {
                "fn_curve": "社会风险 F-N 曲线，展示死亡人数阈值与累积频率的关系。",
                "segment_pll_ranking": "管段 PLL 排序柱状图，展示前十个管段的风险贡献。",
                "risk_matrix": "风险矩阵展示图，说明严重度与频率等级的组合分区。",
            }[chart_id]
            _add_docx_picture(
                paragraph.add_run(),
                charts[chart_id],
                width=Inches(6.25),
                alt_text=alt_text,
            )
            caption = document.add_paragraph(
                {
                    "fn_curve": "图：社会风险 F-N 曲线",
                    "segment_pll_ranking": "图：管段 PLL 排序",
                    "risk_matrix": "图：风险矩阵展示",
                }[chart_id]
            )
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption.runs:
                _set_docx_font(run)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(102, 120, 138)

    document.add_heading("附录 A  节点状态目录", level=1)
    _docx_table(
        document,
        [
            ["序号", "节点", "状态", "依据"],
            *[
                [
                    str(row["sequence_no"]),
                    str(row["label_zh"] or row["node_id"]),
                    str(row["status"]),
                    str(row["standard_ref"] or "—"),
                ]
                for row in context["nodes"]
            ],
        ],
        [700, 2700, 1600, 4360],
    )
    document.add_heading("附录 B  管段风险排序", level=1)
    _docx_table(
        document,
        [
            ["排序", "管段", "PLL", "占比"],
            *[
                [
                    str(row.get("risk_value_rank") or index),
                    str(row.get("segment_id") or "—"),
                    _scientific(row.get("risk_value_fatalities_per_year")),
                    _percent(row.get("fraction_of_pipeline_risk_value")),
                ]
                for index, row in enumerate(context["segment_ranking"][:10], start=1)
            ],
        ],
        [1000, 1800, 3500, 3060],
    )
    document.add_heading("附录 C  F-N 结果点", level=1)
    _docx_table(
        document,
        [
            ["死亡人数下限", "累积频率"],
            *[
                [
                    _decimal(row.get("fatalities_at_least"), 1),
                    _scientific(row.get("cumulative_frequency_per_year")),
                ]
                for row in context["fn_curve"]
            ],
        ],
        [3600, 5760],
    )
    document.add_heading("附录 D  正式发布阻断项", level=1)
    for blocker in context["formal_release_blockers"]:
        paragraph = document.add_paragraph(style="Normal")
        run = paragraph.add_run(f"阻断项：{blocker}")
        _set_docx_font(run)
    document.add_heading("附录 E  证据索引", level=1)
    _docx_table(
        document,
        [
            ["证据ID", "名称", "SHA-256"],
            *[
                [str(row["evidence_id"]), str(row["label"]), str(row["sha256"])]
                for row in context["evidence_index"]
            ],
        ],
        [2300, 2500, 4560],
    )
    document.add_heading("附录 F  结果索引", level=1)
    _docx_table(
        document,
        [
            ["结果引用", "状态", "SHA-256"],
            *[
                [str(row["result_ref"]), str(row["status"]), str(row["sha256"])]
                for row in context["result_index"]
            ],
        ],
        [2400, 1600, 5360],
    )
    document.add_heading("附录 G  受控哈希", level=1)
    for label, value in (
        ("上下文哈希", json_sha256(context)),
        ("草稿哈希", json_sha256(draft)),
        ("校验状态", validation["status"]),
    ):
        paragraph = document.add_paragraph()
        label_run = paragraph.add_run(f"{label}：")
        _set_docx_font(label_run)
        label_run.bold = True
        value_run = paragraph.add_run(str(value))
        _set_docx_font(value_run)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _render_validation_issues(
    context: dict[str, Any],
    draft: dict[str, Any],
    html_content: bytes,
) -> list[str]:
    issues: list[str] = []
    text = html_content.decode("utf-8")
    if "{{metric:" in text:
        issues.append("HTML中仍有未替换的指标占位符")
    for section in draft["sections"]:
        for paragraph in section["paragraphs"]:
            rendered = render_text_template(paragraph["text_template"], context)
            if escape(rendered) not in text:
                issues.append(f"HTML缺少受控段落：{paragraph['paragraph_id']}")
    if SYNTHETIC_WATERMARK not in text:
        issues.append("HTML缺少合成数据水印")
    return issues


class ControlledReportService:
    def __init__(self, database: QraDatabase):
        self.database = database

    def generate(
        self,
        project_id: str,
        *,
        provider: StructuredReportProvider | None = None,
        actor: str = "local-user",
    ) -> ControlledReportBuild:
        context, charts = build_report_context(self.database, project_id)
        fallback_reason: str | None = None
        if provider is None:
            draft = build_deterministic_draft(context)
        else:
            try:
                candidate = provider.generate(
                    prompt=PROMPT_PATH.read_text(encoding="utf-8"),
                    context=deepcopy(context),
                )
                if not isinstance(candidate, dict):
                    raise ValueError("结构化报告模型必须返回JSON对象")
                draft = candidate
                preflight = validate_controlled_report(context, draft)
                if preflight["status"] != "PASS":
                    raise ValueError("模型草稿未通过受控引用和表述校验")
            except Exception as exc:
                fallback_reason = f"{type(exc).__name__}: {str(exc)[:240]}"
                draft = build_deterministic_draft(context)

        validation = validate_controlled_report(context, draft)
        if fallback_reason:
            validation["generation_fallback"] = {
                "requested_provider_id": str(getattr(provider, "provider_id", "unknown")),
                "reason": fallback_reason,
            }
        if validation["status"] != "PASS":
            raise ValueError("确定性报告校验未通过，不能生成报告产物")
        html_content = render_controlled_html(context, draft, validation, charts)
        render_issues = _render_validation_issues(context, draft, html_content)
        validation["checks"].append(_check("RENDERED_NUMERIC_AND_WATERMARK", render_issues))
        validation["status"] = (
            "PASS" if all(row["status"] == "PASS" for row in validation["checks"]) else "FAIL"
        )
        if validation["status"] != "PASS":
            raise ValueError("报告排版后的数字或水印校验未通过")
        pdf_content = render_controlled_pdf(context, draft, validation, charts)
        docx_content = render_controlled_docx(context, draft, validation, charts)
        if not pdf_content.startswith(b"%PDF"):
            raise ValueError("PDF产物签名无效")
        if not docx_content.startswith(b"PK"):
            raise ValueError("DOCX产物签名无效")
        report = self.database.create_controlled_report(
            project_id=project_id,
            run_id=str(context["calculation_run"]["id"]),
            generation_mode=str(draft["generation_mode"]),
            provider_id=str(draft["provider_id"]),
            prompt_version=REPORT_PROMPT_VERSION,
            context=context,
            draft=draft,
            validation=validation,
            html_content=html_content,
            pdf_content=pdf_content,
            docx_content=docx_content,
            actor=actor,
        )
        return ControlledReportBuild(
            report=report,
            context=context,
            draft=draft,
            validation=validation,
            html=html_content,
            pdf=pdf_content,
            docx=docx_content,
            charts=charts,
        )

    def confirm(
        self,
        report_id: str,
        *,
        reviewer: str,
        reason: str,
    ) -> dict[str, Any]:
        integrity = self.verify_integrity(report_id)
        if integrity["status"] != "PASS":
            self.database.record_event(
                event_type="CONTROLLED_REPORT_CONFIRMATION_REJECTED",
                entity_type="controlled_report",
                entity_id=report_id,
                detail={
                    "reason": "integrity_validation_failed",
                    "failed_checks": [
                        row["check_id"]
                        for row in integrity["checks"]
                        if row["status"] != "PASS"
                    ],
                },
                actor=str(reviewer).strip()[:120] or "local-reviewer",
            )
            raise ValueError("报告完整性或数字引用校验失败，不能人工确认")
        try:
            return self.database.confirm_controlled_report(
                report_id,
                reviewer=reviewer,
                reason=reason,
            )
        except ValueError as exc:
            self.database.record_event(
                event_type="CONTROLLED_REPORT_CONFIRMATION_REJECTED",
                entity_type="controlled_report",
                entity_id=report_id,
                detail={"reason": str(exc)[:240]},
                actor=str(reviewer).strip()[:120] or "local-reviewer",
            )
            raise

    def validate_confirmation_candidate(
        self,
        report_id: str,
        draft: dict[str, Any],
        *,
        actor: str = "local-reviewer",
    ) -> dict[str, Any]:
        """Validate a review-time draft without ever mutating the stored report version."""

        report = self.database.get_controlled_report(report_id)
        validation = validate_controlled_report(report["context"], draft)
        draft_matches = json_sha256(draft) == str(report["draft_sha256"])
        validation["checks"].append(
            _check(
                "CONFIRMATION_DRAFT_MATCH",
                [] if draft_matches else ["复核草稿与不可变报告版本不一致"],
            )
        )
        validation["status"] = (
            "PASS"
            if all(row["status"] == "PASS" for row in validation["checks"])
            else "FAIL"
        )
        if validation["status"] != "PASS":
            self.database.record_event(
                event_type="CONTROLLED_REPORT_CONFIRMATION_REJECTED",
                entity_type="controlled_report",
                entity_id=report_id,
                detail={
                    "reason": "candidate_validation_failed",
                    "candidate_draft_sha256": json_sha256(draft),
                    "failed_checks": [
                        row["check_id"]
                        for row in validation["checks"]
                        if row["status"] != "PASS"
                    ],
                },
                actor=str(actor).strip()[:120] or "local-reviewer",
            )
        return validation

    def verify_integrity(self, report_id: str) -> dict[str, Any]:
        """Recompute all stored report hashes and report rules before confirmation."""

        report = self.database.get_controlled_report(report_id)
        validation = validate_controlled_report(report["context"], report["draft"])
        checks = list(validation["checks"])
        hash_issues: list[str] = []
        if json_sha256(report["context"]) != str(report["context_sha256"]):
            hash_issues.append("报告上下文哈希不一致")
        if json_sha256(report["draft"]) != str(report["draft_sha256"]):
            hash_issues.append("报告草稿哈希不一致")
        html_content = b""
        for format_name in ("html", "pdf", "docx"):
            _, content, _ = self.database.get_controlled_report_artifact(report_id, format_name)
            if bytes_sha256(content) != str(report[f"{format_name}_sha256"]):
                hash_issues.append(f"{format_name.upper()}产物哈希不一致")
            if format_name == "html":
                html_content = content
        checks.append(_check("STORED_REPORT_HASHES", hash_issues))
        checks.append(
            _check(
                "STORED_RENDERED_NUMERIC_AND_WATERMARK",
                _render_validation_issues(report["context"], report["draft"], html_content),
            )
        )
        return {
            "schema_version": REPORT_VALIDATION_VERSION,
            "status": (
                "PASS" if all(row["status"] == "PASS" for row in checks) else "FAIL"
            ),
            "checks": checks,
        }


def build_controlled_report_zip(database: QraDatabase, report_id: str) -> bytes:
    report = database.get_controlled_report(report_id)
    output = io.BytesIO()
    html_type, html_content, _ = database.get_controlled_report_artifact(report_id, "html")
    _, pdf_content, _ = database.get_controlled_report_artifact(report_id, "pdf")
    _, docx_content, _ = database.get_controlled_report_artifact(report_id, "docx")
    chart_matches = re.findall(
        rb"data:image/png;base64,([A-Za-z0-9+/=]+)",
        html_content,
    )
    chart_resources = report["context"].get("chart_resources", [])
    embedded_charts = {
        bytes_sha256(content): content
        for content in (base64.b64decode(encoded) for encoded in chart_matches)
    }
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("report.html", html_content)
        archive.writestr("report.pdf", pdf_content)
        archive.writestr("report.docx", docx_content)
        archive.writestr(
            "report-context-v1.json",
            json.dumps(
                report["context"], ensure_ascii=False, indent=2, sort_keys=True, allow_nan=False
            ).encode("utf-8"),
        )
        archive.writestr(
            "report-draft-v1.json",
            json.dumps(
                report["draft"], ensure_ascii=False, indent=2, sort_keys=True, allow_nan=False
            ).encode("utf-8"),
        )
        archive.writestr(
            "report-validation-v1.json",
            json.dumps(
                report["validation"],
                ensure_ascii=False,
                indent=2,
                sort_keys=True,
                allow_nan=False,
            ).encode("utf-8"),
        )
        for resource in chart_resources:
            content = embedded_charts.get(str(resource["sha256"]))
            if content is None:
                raise ValueError(f"报告图表哈希不一致：{resource['chart_id']}")
            archive.writestr(str(resource["path"]), content)
        manifest = {
            "schema_version": "controlled-report-bundle-v1",
            "report_id": report_id,
            "status": report["status"],
            "context_sha256": report["context_sha256"],
            "draft_sha256": report["draft_sha256"],
            "html": {"content_type": html_type, "sha256": report["html_sha256"]},
            "pdf_sha256": report["pdf_sha256"],
            "docx_sha256": report["docx_sha256"],
            "formal_report_allowed": bool(report["context"]["formal_report_allowed"]),
        }
        archive.writestr(
            "bundle-manifest.json",
            json.dumps(
                manifest, ensure_ascii=False, indent=2, sort_keys=True, allow_nan=False
            ).encode("utf-8"),
        )
    return output.getvalue()


__all__ = [
    "ControlledReportBuild",
    "ControlledReportService",
    "FALLBACK_PROVIDER_ID",
    "REPORT_CONTEXT_VERSION",
    "REPORT_DRAFT_VERSION",
    "REPORT_PROMPT_VERSION",
    "REPORT_SECTIONS",
    "REPORT_SERVICE_VERSION",
    "SYNTHETIC_WATERMARK",
    "StructuredReportProvider",
    "build_chart_bundle",
    "build_controlled_report_zip",
    "build_deterministic_draft",
    "build_report_context",
    "render_controlled_docx",
    "render_controlled_html",
    "render_controlled_pdf",
    "render_text_template",
    "validate_controlled_report",
]
